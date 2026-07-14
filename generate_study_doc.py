from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(214, 40, 40)

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = RGBColor(80, 80, 80)

def p(text, bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.bold = bold
    return para

def bullet(text):
    para = doc.add_paragraph(text, style='List Bullet')
    return para

def code_block(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return para

def make_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D62828"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return table


# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p('BloodBridge', bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Digital Blood Bank System', bold=True)
p('Complete Project Explanation Document')
p('')
p('Course: CSE 3100 - Web Based Application Project')
p('Team: Tanzirul Islam | Yeanur Hossain | Md. Emon Islam')
p('Supervisor: Md. Rabiul Islam, Professor, Dept. of CSE, RUET')
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. WHAT IS BLOODBIDGE
# ════════════════════════════════════════════════════════════
h1('1. BloodBridge Ki? (Project Overview)')

p('BloodBridge ekta web-based digital blood bank system jeta blood donor der ke patient/recipient der sathe connect kore. '
  'Bangladesh e onek somoy emergency e blood lagle social media e post korte hoy, phone call korte hoy - '
  'eta organized na. BloodBridge ei problem solve kore ekta centralized platform diye.')

p('简而言之 (Simply put): Ekta website jekhane:')
bullet('Donor register korte pare blood group + location shoho')
bullet('Kono manush blood request post korte pare (ki blood lagche, kothay, kokhono)')
bullet('Donor search korte pare blood group + district + upazila diye')
bullet('Admin sob manage korte pare (users, requests, blogs)')
bullet('Blog module ache health awareness content er jonno')

h2('1.1 Ki ki Feature Ache')
make_table(['Feature', 'Description'], [
    ['User Registration', 'Firebase auth diye register, blood group + location info store hoy MongoDB e'],
    ['Login', 'Email/password ba Google sign-in'],
    ['Blood Request Post', 'Recipient name, blood group, hospital, district, date shoho create korte paro'],
    ['Donate to Request', 'Kono pending request e donor "I will donate" click korte paro'],
    ['Donor Search', 'Blood group, district, upazila diye filter kore donor khujhte paro'],
    ['Blog Module', 'Admin/Volunteer blog lekhe publish korte paro'],
    ['Admin Dashboard', 'Stats dekhte pare, users manage korte pare, requests manage korte pare'],
    ['Role-based Access', '3 ta role: Donor, Volunteer, Admin - pratekar alada access'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. TECH STACK
# ════════════════════════════════════════════════════════════
h1('2. Tech Stack - Ki Ki Use Korsi & Keno')

p('BloodBridge MERN stack use kore build koreche. MERN mane MongoDB + Express.js + React + Node.js.')

h2('2.1 Frontend (Client Side)')
make_table(['Technology', 'Version', 'Ki Jonno', 'Keno Eta Choose Korsi'], [
    ['React', '18.3.1', 'UI build', 'Component-based, reusable UI banana easy, industry standard'],
    ['Vite', '5.4.0', 'Dev server + build tool', 'CRA er theke 10x fast, instant hot reload'],
    ['Tailwind CSS', '4.2.2', 'Styling', 'Utility classes diye quickly design kora jay, custom CSS lage na'],
    ['Firebase Auth', '12.12.0', 'Authentication', 'Google sign-in + email/password free e paoa jay, secure'],
    ['Axios', '1.15.0', 'HTTP requests', 'Backend API call kora, interceptor token auto-attach kore'],
    ['React Router', '7.14.1', 'Page routing', 'URL e /dashboard, /login etc handle kore SPA te'],
    ['React Hot Toast', '2.6.0', 'Notifications', 'Success/error messages popup e dekhay'],
    ['SweetAlert2', '11.26.24', 'Confirmation dialogs', 'Delete korar age "Are you sure?" modal'],
    ['React Quill', '2.0.0', 'Rich text editor', 'Blog post e formatted text likhte use hoy'],
])

h2('2.2 Backend (Server Side)')
make_table(['Technology', 'Version', 'Ki Jonno', 'Keno Eta Choose Korsi'], [
    ['Node.js', '22.x', 'JavaScript runtime', 'Server-side JS run kore, same language frontend-backend'],
    ['Express.js', '4.18.2', 'Web framework', 'REST API banano simple, minimal boilerplate'],
    ['MongoDB + Mongoose', '8.0.3', 'Database', 'Flexible schema, JSON-like documents, free cloud hosting'],
    ['jsonwebtoken', '9.0.2', 'JWT tokens', 'Stateless auth, lightweight, industry standard'],
    ['cors', '2.8.5', 'Cross-origin requests', 'Frontend different port theke API call korte pare'],
    ['dotenv', '16.3.1', 'Environment variables', '.env file theke secrets load kore, security'],
])

h2('2.3 Hosting & Tools')
make_table(['Tool', 'Purpose'], [
    ['MongoDB Atlas', 'Free cloud database (M0 cluster, 512MB)'],
    ['Firebase', 'Free authentication (Spark plan)'],
    ['imgbb API', 'Free image hosting for profile photos and blog thumbnails'],
    ['Vercel', 'Frontend + Backend deployment'],
    ['GitHub', 'Version control + collaboration'],
    ['VS Code', 'Code editor'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. FOLDER STRUCTURE
# ════════════════════════════════════════════════════════════
h1('3. Folder Structure - Project Ki Vabe Organized')

p('Project ta 2 ta major folder e divided:')

h2('3.1 Complete Folder Tree')
code = """BloodBridge/
├── BloodBridge-client/          <-- Frontend (React app)
│   ├── src/
│   │   ├── components/          <-- Reusable UI parts
│   │   │   ├── Navbar.jsx       <-- Top navigation bar
│   │   │   ├── Footer.jsx       <-- Bottom footer
│   │   │   ├── PrivateRoute.jsx <-- Login required pages protect kore
│   │   │   ├── DashboardLayout.jsx <-- Dashboard sidebar + content area
│   │   │   └── LoadingSpinner.jsx  <-- Loading animation
│   │   ├── pages/               <-- Full page components
│   │   │   ├── Home.jsx         <-- Landing page
│   │   │   ├── Login.jsx        <-- Login form
│   │   │   ├── Register.jsx     <-- Registration form
│   │   │   ├── AllRequests.jsx  <-- Public blood requests list
│   │   │   ├── RequestDetail.jsx <-- Single request er details
│   │   │   ├── SearchDonors.jsx <-- Donor search page
│   │   │   ├── Blog.jsx         <-- Blog listing
│   │   │   ├── BlogDetail.jsx   <-- Single blog post
│   │   │   ├── NotFound.jsx     <-- 404 page
│   │   │   └── dashboard/       <-- Dashboard pages
│   │   │       ├── DashboardHome.jsx
│   │   │       ├── Profile.jsx
│   │   │       ├── CreateRequest.jsx
│   │   │       ├── MyRequests.jsx
│   │   │       ├── EditRequest.jsx
│   │   │       ├── AllUsers.jsx         <-- Admin only
│   │   │       ├── AllRequestsAdmin.jsx <-- Admin only
│   │   │       ├── ContentManagement.jsx
│   │   │       └── AddBlog.jsx
│   │   ├── context/
│   │   │   └── AuthContext.jsx  <-- Global auth state (login/logout/register)
│   │   ├── firebase/
│   │   │   └── firebase.config.js <-- Firebase setup
│   │   ├── hooks/
│   │   │   └── useAxios.js      <-- Axios instance with token
│   │   ├── utils/
│   │   │   └── imgbbUpload.js   <-- Image upload helper
│   │   ├── data/
│   │   │   └── bangladesh.js    <-- 64 districts + upazilas data
│   │   ├── main.jsx             <-- App entry point, routing
│   │   └── style.css            <-- Global styles, Tailwind theme
│   ├── package.json
│   ├── vite.config.js
│   ├── tailwind.config.js
│   └── .env                     <-- Firebase keys, API URL
│
├── BloodBridge-server/          <-- Backend (Express API)
│   ├── api/
│   │   └── index.js             <-- Vercel serverless entry (skip for now)
│   ├── models/                  <-- MongoDB schemas
│   │   ├── User.js              <-- User data structure
│   │   ├── DonationRequest.js   <-- Blood request data structure
│   │   └── Blog.js              <-- Blog post data structure
│   ├── routes/                  <-- API endpoints
│   │   ├── auth.js              <-- /api/auth/* (register, login, verify)
│   │   ├── users.js             <-- /api/users/* (CRUD, search, admin)
│   │   ├── donationRequests.js  <-- /api/donationRequests/* (CRUD, donate, complete)
│   │   ├── blogs.js             <-- /api/blogs/* (CRUD, publish/unpublish)
│   │   └── stats.js             <-- /api/stats (admin dashboard numbers)
│   ├── middleware/
│   │   └── verifyToken.js       <-- JWT token check middleware
│   ├── server.js                <-- Main server file (Express app setup)
│   ├── package.json
│   └── .env                     <-- MongoDB URI, JWT secret"""
code_block(code)

h2('3.2 Key Files Explanation')
make_table(['File', 'Eta Ki Kore'], [
    ['main.jsx', 'React app start hoy ekhane. BrowserRouter + AuthProvider wrap kore pura app ke.'],
    ['AuthContext.jsx', 'Ekta React Context jeta global e store kore: user login ache kina, dbUser er data, login/logout/register functions.'],
    ['server.js', 'Express app create kore, routes register kore, MongoDB connect kore, port e listen kore.'],
    ['verifyToken.js', 'Prottek protected request e token check kore. Token valid hole req.user e decoded data attach kore.'],
    ['bangladesh.js', '64 ta district er list + pratekar upazila gula. Dropdown populate korte use hoy.'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. DATABASE DESIGN
# ════════════════════════════════════════════════════════════
h1('4. Database Design - MongoDB Schema')

p('MongoDB Atlas e 3 ta collection ache:')

h2('4.1 users Collection')
make_table(['Field', 'Type', 'Ki Store Hoy'], [
    ['_id', 'ObjectId', 'Auto-generated unique ID'],
    ['name', 'String', 'User er full name'],
    ['email', 'String', 'Email address (unique)'],
    ['avatar', 'String', 'Profile photo URL (imgbb hosted)'],
    ['role', 'String', 'admin / volunteer / donor (default: donor)'],
    ['bloodGroup', 'String', 'A+, A-, B+, B-, AB+, AB-, O+, O-'],
    ['district', 'String', 'Bangladesh district name'],
    ['upazila', 'String', 'Sub-district name'],
    ['status', 'String', 'active / blocked (default: active)'],
    ['createdAt', 'Date', 'Auto-generated timestamp'],
])

h2('4.2 donationrequests Collection')
make_table(['Field', 'Type', 'Ki Store Hoy'], [
    ['_id', 'ObjectId', 'Auto-generated unique ID'],
    ['requesterName', 'String', 'Jei person request ta create korse'],
    ['requesterEmail', 'String', 'Tar email'],
    ['recipientName', 'String', 'Jake blood lagche'],
    ['bloodGroup', 'String', 'Ki blood group lagche'],
    ['hospitalName', 'String', 'Kon hospital e'],
    ['fullAddress', 'String', 'Full address'],
    ['district', 'String', 'District'],
    ['upazila', 'String', 'Upazila'],
    ['donationDate', 'Date', 'Kokhon lagbe'],
    ['donationTime', 'String', 'Ki somoy lagbe'],
    ['requestMessage', 'String', 'Optional message'],
    ['status', 'String', 'pending / inprogress / done / canceled'],
    ['donorInfo', 'Object', '{ name, email } - ke donate korbe'],
])

h2('4.3 blogs Collection')
make_table(['Field', 'Type', 'Ki Store Hoy'], [
    ['_id', 'ObjectId', 'Auto-generated unique ID'],
    ['title', 'String', 'Blog title'],
    ['thumbnail', 'String', 'Thumbnail image URL'],
    ['content', 'String', 'Blog body (HTML from React Quill)'],
    ['status', 'String', 'draft / published'],
    ['authorName', 'String', 'Author er name'],
    ['authorEmail', 'String', 'Author er email'],
])

h2('4.4 Request Status Flow')
p('Donation Request er lifecycle:')
code_block("""pending  -->  inprogress  -->  done
   |                              |
   +----> canceled <--------------+""")
p('pending: Request create hoise, kono donor accept kore nai.')
p('inprogress: Ektu donor "Donate" e click korse, donate korbe.')
p('done: Donate hoye geche.')
p('canceled: Request cancel hoye geche.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. AUTHENTICATION FLOW
# ════════════════════════════════════════════════════════════
h1('5. Authentication Flow - Kivabe Login/Register Kaj Kore')

p('Authentication 2 step e kaj kore: Firebase + JWT.')

h2('5.1 Registration Flow')
code_block("""User (Browser)                    Firebase                    Backend API
    |                                  |                           |
    |-- Register form fill up -------->|                           |
    |-- createUserWithEmailAndPassword->|                           |
    |<-- Firebase user created --------|                           |
    |-- Send user data to POST /register ------------------------->|
    |                                  |                           |-- Save to MongoDB
    |                                  |                           |-- Generate JWT
    |<-- JWT token + user data -------|                           |
    |-- Store token in localStorage ---|                           |""")
p('Short e: Firebase e user create hoy -> Backend e data save hoy -> JWT token pawa jay.')

h2('5.2 Login Flow')
code_block("""User (Browser)                    Firebase                    Backend API
    |                                  |                           |
    |-- Email + Password submit ------->|                           |
    |-- signInWithEmailAndPassword ---->|                           |
    |<-- Firebase ID Token ------------|                           |
    |-- Send email to POST /jwt ------>|-------------------------->|
    |                                  |                           |-- Find user in MongoDB
    |                                  |                           |-- Generate JWT (7 days)
    |<-- JWT token + user data -------|<---------------------------|
    |-- Store token in localStorage    |                           |""")

h2('5.3 How Protected Routes Work')
p('1. User login korle JWT token localStorage e store hoy.')
p('2. Prottek API request e Axios interceptor token ta Authorization header e attach kore.')
p('3. Backend e verifyToken middleware token verify kore.')
p('4. Token valid hole request age jay, na hole 401 return hoy.')
p('5. Frontend e PrivateRoute component check kore: user ache kina. Na hole /login e redirect.')

h2('5.4 Three User Roles')
make_table(['Role', 'Ki Korte Pare', 'Dashboard e Ki Dekhe'], [
    ['Donor', 'Register, create request, donate, edit/delete own requests', 'Own requests, profile, create request'],
    ['Volunteer', 'Create/manage blog posts, view dashboard', 'Welcome message, content management'],
    ['Admin', 'Everything: manage users, requests, blogs, view stats', 'Stats cards, all users, all requests, content management'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. API ENDPOINTS
# ════════════════════════════════════════════════════════════
h1('6. API Endpoints - Backend Routes Gula')

h2('6.1 Authentication (/api/auth)')
make_table(['Method', 'Endpoint', 'Auth?', 'Ki Kore'], [
    ['POST', '/api/auth/jwt', 'No', 'Email diye login, JWT token return kore'],
    ['POST', '/api/auth/register', 'No', 'Notun user register kore'],
    ['GET', '/api/auth/verify', 'Yes', 'Token verify kore, current user return kore'],
])

h2('6.2 Users (/api/users)')
make_table(['Method', 'Endpoint', 'Auth?', 'Ki Kore'], [
    ['GET', '/api/users/', 'Admin', 'Shob user list (paginated)'],
    ['GET', '/api/users/donors', 'No', 'Donor search (bloodGroup, district, upazila filter)'],
    ['PUT', '/api/users/:id', 'Yes', 'User profile update'],
    ['PUT', '/api/users/:id/block', 'Admin', 'User ke block/unblock'],
    ['PUT', '/api/users/:id/role', 'Admin', 'User er role change'],
    ['DELETE', '/api/users/:id', 'Admin', 'User delete'],
])

h2('6.3 Donation Requests (/api/donationRequests)')
make_table(['Method', 'Endpoint', 'Auth?', 'Ki Kore'], [
    ['GET', '/api/donationRequests/', 'No', 'Shob requests (paginated, filterable)'],
    ['GET', '/api/donationRequests/pending', 'No', 'Shobar recent 3 ta pending request (home page e)'],
    ['GET', '/api/donationRequests/:id', 'No', 'Single request er details'],
    ['POST', '/api/donationRequests/', 'Yes', 'Notun blood request create'],
    ['PUT', '/api/donationRequests/:id', 'Yes', 'Request edit'],
    ['PUT', '/api/donationRequests/:id/donate', 'Yes', 'Request e donate korar promise (status -> inprogress)'],
    ['PUT', '/api/donationRequests/:id/complete', 'Yes', 'Request done标记'],
    ['PUT', '/api/donationRequests/:id/cancel', 'Yes', 'Request cancel'],
    ['DELETE', '/api/donationRequests/:id', 'Yes', 'Request delete'],
])

h2('6.4 Blogs (/api/blogs)')
make_table(['Method', 'Endpoint', 'Auth?', 'Ki Kore'], [
    ['GET', '/api/blogs/', 'No', 'Shob published blogs'],
    ['GET', '/api/blogs/all', 'Yes', 'Shob blogs (draft included)'],
    ['POST', '/api/blogs/', 'Yes', 'Notun blog create (draft)'],
    ['PUT', '/api/blogs/:id/publish', 'Yes', 'Blog publish'],
    ['PUT', '/api/blogs/:id/unpublish', 'Yes', 'Blog draft e felo'],
    ['DELETE', '/api/blogs/:id', 'Yes', 'Blog delete'],
])

h2('6.5 Stats (/api/stats)')
make_table(['Method', 'Endpoint', 'Auth?', 'Ki Kore'], [
    ['GET', '/api/stats/', 'Admin', 'Dashboard stats: totalDonors, totalVolunteers, totalRequests, completed'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. FRONTEND PAGES & ROUTING
# ════════════════════════════════════════════════════════════
h1('7. Frontend Pages & Routing')

h2('7.1 Public Routes (Navbar + Footer show hoy)')
make_table(['URL', 'Page', 'Ki Kore'], [
    ['/', 'Home', 'Hero section, how it works, urgent requests, stats, CTA buttons'],
    ['/login', 'Login', 'Email/password + Google login form (Navbar hidden)'],
    ['/register', 'Register', 'Full registration form (Navbar hidden)'],
    ['/blood-donation-requests', 'AllRequests', 'Shob pending requests card e show kore'],
    ['/blood-donation-requests/:id', 'RequestDetail', 'Single request er details + Donate button'],
    ['/search-donors', 'SearchDonors', 'Blood group + district + upazila filter diye donor search'],
    ['/blog', 'Blog', 'Published blogs er listing'],
    ['/blog/:id', 'BlogDetail', 'Single blog er full content'],
])

h2('7.2 Protected Routes (Login Lagbe)')
make_table(['URL', 'Page', 'Roles', 'Ki Kore'], [
    ['/dashboard', 'DashboardHome', 'All', 'Admin: stats cards. Donor: recent requests. Volunteer: welcome.'],
    ['/dashboard/profile', 'Profile', 'All', 'Name, blood group, district, upazila, avatar edit'],
    ['/dashboard/create-donation-request', 'CreateRequest', 'Donor', 'Notun blood request form'],
    ['/dashboard/my-donation-requests', 'MyRequests', 'Donor', 'Own requests table + edit/delete/done/cancel actions'],
    ['/dashboard/edit-donation-request/:id', 'EditRequest', 'Donor', 'Existing request edit form'],
    ['/dashboard/all-users', 'AllUsers', 'Admin', 'Paginated user table, role change, block/unblock, delete'],
    ['/dashboard/all-blood-donation-request', 'AllRequestsAdmin', 'Admin', 'Paginated all requests table with status filter'],
    ['/dashboard/content-management', 'ContentManagement', 'Admin,Vol', 'Blog manage: publish/unpublish/delete'],
    ['/dashboard/add-blog', 'AddBlog', 'Admin,Vol', 'Notun blog create form with rich text editor'],
])

h2('7.3 How Routing Works')
p('React Router DOM use kore routing. main.jsx te Routes define kora ache.')
p('Key concept: PrivateRoute component check kore user logged in kina. Na hole /login e redirect kore.')
p('DashboardLayout ekta wrapper jeta sidebar + Outlet (child page) render kore.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. HOW IT ALL CONNECTS
# ════════════════════════════════════════════════════════════
h1('8. How Everything Connects - Full Data Flow')

h2('8.1 Registration Example (End to End)')
code_block("""Step 1: User /register page e form fill up kore
        - Name, email, blood group, district, upazila, photo, password

Step 2: Register.jsx handleSubmit() call hoy
        - Password validate kore (length, uppercase, lowercase, match)

Step 3: AuthContext.register() call hoy
        - Firebase e createUserWithEmailAndPassword() call hoy
        - Firebase user create hoy, ID token pawa jay

Step 4: Backend e POST /api/auth/register call hoy
        - Body te: name, email, avatar, bloodGroup, district, upazila
        - auth.js route check kore: email already exist kina
        - Na thakle User model e notun document create hoy MongoDB te
        - JWT token generate hoy (payload: id, email, role; expiry: 7 days)
        - Response: { token, user: { id, name, email, role, avatar } }

Step 5: Client side e
        - Token localStorage e store hoy
        - dbUser state update hoy
        - User ke home page e redirect kora hoy
        - Toast show hoy: "Registration successful!" """)

h2('8.2 Blood Request Example (End to End)')
code_block("""Step 1: Donor logs in -> Dashboard -> "Create Request" page e jay

Step 2: Form fill up: recipient name, blood group, hospital, address, district, 
        upazila, date, time, optional message

Step 3: CreateRequest.jsx handleSubmit() call hoy

Step 4: Axios POST /api/donationRequests call hoy
        - Header e: Authorization: Bearer <JWT token>
        - Body te: shob form data

Step 5: Backend e verifyToken middleware token verify kore
        - Token valid: req.user e decoded data attach hoy
        - Token invalid: 401 return hoy

Step 6: donationRequests.js route handler
        - Notun DonationRequest document create hoy MongoDB te
        - Status "pending" set hoy
        - Response: { message, request }

Step 7: Client side e
        - My Requests page e redirect hoy
        - Toast: "Request created successfully!"
        - Home page e urgent requests section e ei request dekhte paben""")

h2('8.3 Donor Search Example')
code_block("""Step 1: /search-donors page e filter select kore
        - Blood Group: A+
        - District: Dhaka
        - Upazila: Dhanmondi (auto-loaded district select korar por)

Step 2: SearchDonors.jsx useEffect trigger hoy
        - GET /api/users/donors?bloodGroup=A%2B&district=Dhaka&upazila=Dhanmondi

Step 3: Backend e users.js route
        - MongoDB query: User.find({ role: 'donor', status: 'active', bloodGroup, district, upazila })
        - Matching donors return hoy (name, email, avatar, bloodGroup, location)

Step 4: Results cards e show hoy
        - Each card: avatar, name, blood group badge, location, email""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9. KEY CODE SNIPPETS
# ════════════════════════════════════════════════════════════
h1('9. Key Code Snippets - Important Parts')

h2('9.1 Server Entry (server.js)')
code_block("""// server.js - Express app setup
const express = require('express');
const mongoose = require('mongoose');
const cors = require('cors');

const app = express();
app.use(cors());          // Allow all origins
app.use(express.json());  // Parse JSON bodies

// Register route handlers
app.use('/api/auth', authRoutes);
app.use('/api/users', userRoutes);
// ... more routes

// Connect MongoDB then start server
mongoose.connect(process.env.MONGODB_URI)
  .then(() => app.listen(PORT))
  .catch(err => console.log(err));""")

h2('9.2 JWT Middleware (verifyToken.js)')
code_block("""// verifyToken.js - Check if request has valid JWT
const jwt = require('jsonwebtoken');

const verifyToken = async (req, res, next) => {
  const token = req.headers.authorization?.split(" ")[1];
  // Token header e thake: "Bearer eyJhbGciOi..."
  // Split diye "eyJhbGciOi..." part ta nibo

  if (!token) return res.status(401).json({ message: "unauthorized" });

  const decoded = await jwt.verify(token, process.env.JWT_SECRET);
  req.user = decoded;  // { id, email, role }
  next();
};""")

h2('9.3 Auth Context (AuthContext.jsx)')
code_block("""// AuthContext.jsx - Global auth state management
const AuthProvider = ({ children }) => {
  const [user, setUser] = useState(null);       // Firebase user
  const [dbUser, setDbUser] = useState(null);    // MongoDB user
  const [loading, setLoading] = useState(true);

  const login = async (email, password) => {
    // Step 1: Firebase e login
    const result = await signInWithEmailAndPassword(auth, email, password);
    // Step 2: Backend theke JWT token nao
    const response = await axios.post('/api/auth/jwt', { email });
    // Step 3: Token store koro
    localStorage.setItem('token', response.data.token);
    setDbUser(response.data.user);
  };

  return (
    <AuthContext.Provider value={{ user, dbUser, login, register, logout }}>
      {children}
    </AuthContext.Provider>
  );
};""")

h2('9.4 Axios Interceptor (useAxios.js)')
code_block("""// useAxios.js - Auto-attach token to all requests
const useAxios = axios.create({ baseURL: API_BASE_URL });

// Request interceptor: prottek request e token attach koro
useAxios.interceptors.request.use((config) => {
  const token = localStorage.getItem('token');
  if (token) {
    config.headers.Authorization = `Bearer ${token}`;
  }
  return config;
});

// Response interceptor: 404 hole login e redirect koro
useAxios.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem('token');
      window.location.href = '/login';
    }
  }
);""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 10. HOW TO RUN
# ════════════════════════════════════════════════════════════
h1('10. How to Run the Project Locally')

h2('10.1 Prerequisites')
bullet('Node.js installed (v18 or higher)')
bullet('Git installed')
bullet('VS Code editor')

h2('10.2 Setup Steps')
code_block("""Step 1: Clone the repo
  git clone https://github.com/TanzirulIslam22/Blood-Bridge.git
  cd Blood-Bridge

Step 2: Install server dependencies
  cd BloodBridge-server
  npm install

Step 3: Install client dependencies
  cd ../BloodBridge-client
  npm install

Step 4: Setup .env files
  - Server .env (already configured with MongoDB URI + JWT secret)
  - Client .env (already configured with Firebase keys + API URL)

Step 5: Start the server (Terminal 1)
  cd BloodBridge-server
  node server.js
  // Output: "Server running on port 5000" + "MongoDB Connected"

Step 6: Start the client (Terminal 2)
  cd BloodBridge-client
  npm run dev
  // Output: "Local: http://localhost:5173/"

Step 7: Open browser
  http://localhost:5173""")

h2('10.3 Important: Both Must Run Together')
p('CRITICAL: Server ar client 2 ta alada terminal e chalano lobe. '
  'Shudhu client chalaile API call fail hobe. Shudhu server chalaile UI dekhte parben na.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 11. COMMON ISSUES
# ════════════════════════════════════════════════════════════
h1('11. Common Issues & Fixes')

make_table(['Issue', 'Possible Reason', 'Fix'], [
    ['Blank page / white screen', 'Client not running or wrong port', 'Ensure npm run dev is running, check terminal output for port'],
    ['Registration fails', 'Firebase config wrong or server not running', 'Check .env has correct Firebase keys, ensure server is on port 5000'],
    ['401 Unauthorized', 'JWT token expired or not stored', 'Login again, check localStorage has token'],
    ['CORS error in console', 'Client running on different port than 5173', 'Check Vite terminal - it might be on 5174. Kill port 5173 processes'],
    ['MongoDB connection failed', 'Wrong connection string or cluster paused', 'Check .env MONGODB_URI, login to MongoDB Atlas and resume cluster'],
    ['Port 5000 already in use', 'Old server process still running', 'Kill the process: taskkill /F /PID <PID> or restart terminal'],
    ['Request stays "pending"', 'No donor clicked "Donate"', 'Login as another user, go to requests, click Donate'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 12. TEAM CONTRIBUTION
# ════════════════════════════════════════════════════════════
h1('12. Team Contribution & Task Distribution')

make_table(['Member', 'Role', 'Work Done'], [
    ['Tanzirul Islam (2203054)', 'Frontend Lead', 'React components, routing, UI design, Firebase integration, responsive layout, blog module, dashboard pages, Tailwind styling'],
    ['Yeanur Hossain (2203020)', 'Backend Lead', 'Express.js API, MongoDB schemas, JWT auth, route handlers, Vercel deployment, database management, environment config'],
    ['Md. Emon Islam (2203028)', 'Full Stack Support', 'Feature integration, testing, bug fixing, Bangladesh geographic data, content management module, documentation, deployment support'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 13. INTERVIEW QUESTIONS
# ════════════════════════════════════════════════════════════
h1('13. Sir Key Questions & Answers (Viva Prep)')

h2('Q1: MERN stack ki?')
p('MongoDB (database) + Express.js (backend framework) + React (frontend library) + Node.js (JavaScript runtime). '
  'Ei 4 ta mile full-stack web app banano jay sirf JavaScript diye.')

h2('Q2: Firebase + JWT keno 2 ta auth?')
p('Firebase authentication secure ebong easy (Google sign-in, password hashing shob handle kore). '
  'But Firebase token backend e verify kora jay na directly. Tai amra Firebase diye login korai, '
  'then backend theke nijer JWT token niye tai use kore API calls e. '
  'Firebase = identity provider, JWT = API authorization.')

h2('Q3: MongoDB vs SQL database?')
p('MongoDB NoSQL - flexible schema, JSON-like documents, horizontal scaling easy. '
  'BloodBridge er data gula naturally nested (request er shathe donor info, blog er shathe author info) '
  'tai MongoDB suitable. Ektu collection e field add/remove kora simple.')

h2('Q4: REST API ki?')
p('Representational State Transfer - ekta architectural style. Resources ke URLs represent kore '
  '(e.g., /api/users, /api/blogs). HTTP methods (GET, POST, PUT, DELETE) diye operations kore. '
  'Stateless - prottek request independently process hoy.')

h2('Q5: CORS ki? Cross-origin ki?')
p('Cross-Origin Resource Sharing. Jokhon frontend (localhost:5173) backend ke (localhost:5000) call kore, '
  'tara different origin. Browser security rule onujayi, different origin theke request block hoy by default. '
  'CORS middleware diye backend allowlist e specific origins add kore.')

h2('Q6: JWT token ki? Ki thake payload e?')
p('JSON Web Token - ekta signed string. 3 parts: header, payload, signature. '
  'Amader payload e thake: { id, email, role } + expiry (7 days). '
  'Signature ensure kore token ta change hoyni. Backend e verify kore private key diye.')

h2('Q7: Role-based access control (RBAC) ki?')
p('User ke roles assign koro (donor/volunteer/admin), tar onujayi features access kore. '
  'Donor shudhu nijer request manage kore. Admin shob manage kore. '
  'Backend e middleware check kore role, frontend e conditional rendering hoy role onujayi.')

h2('Q8: State management ki use korsi?')
p('React Context API (AuthContext). Redux use kora hoyni - project size small, Context sufficient. '
  'AuthContext global e store kore: user, dbUser, loading, login/logout/register functions. '
  'Component theke useAuth() hook diye access kore.')

h2('Q9: Why Vite not Create React App?')
p('Vite 10x faster than CRA. Hot Module Replacement instant, build time kom. '
  'CRA deprecated hoye geche 2023 theke. Vite lightweight, modern, esbuild based.')

h2('Q10: Deployment kivabe korsi?')
p('Vercel e deploy korechi - client and server 2 ta alada project hisebe. '
  'Client: Vite build output (dist folder) serve kore. Server: Express app ke serverless function e convert kore. '
  'GitHub integration - push korle auto-deploy hoy.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 14. QUICK CHEAT SHEET
# ════════════════════════════════════════════════════════════
h1('14. Quick Cheat Sheet - Sir Ke Bojhano 1-Minute Pitch')

p('Sir, amar project ta BloodBridge - ekta digital blood bank system. '
  'Eta MERN stack e build koreche. Frontend e React + Tailwind CSS, backend e Express.js + MongoDB Atlas. '
  'Firebase authentication use korechi email/password and Google sign-in er jonno.')

p('Project e 3 ta user role ache: donor, volunteer, donor register kore blood group + location diye, '
  'blood request post korte pare. Donor search korte pare blood group, district, upazila diye. '
  'Admin sob manage korte pare - users, requests, blogs. '
  'Blog module ache health awareness content er jonno.')

p('Technical亮点:')
bullet('Firebase Auth + JWT token for secure two-layer authentication')
bullet('Role-based access control - 3 roles with different permissions')
bullet('64 districts + upazilas of Bangladesh for location-based filtering')
bullet('Request lifecycle: pending -> in-progress -> done/canceled')
bullet('Responsive dark-themed UI with Tailwind CSS')
bullet('Deployed on Vercel with GitHub auto-deployment')

p('Team contribution: amar team e 3 jon. Frontend lead, backend lead, aur full-stack support. '
  'Git + GitHub use korechi version control e.')

# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════
output_path = r'D:\All Stack Projects\3100 MERN\BloodBridge\BloodBridge_Study_Document.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
